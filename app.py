# ============================================================
#  Book Layout Generator
#  Streamlit + python-docx
#  Generates a landscape Word doc where each page = 2 book
#  half-pages. Print double-sided, cut vertically, stack.
# ============================================================

import streamlit as st
from docx import Document
from docx.shared import Pt, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import io


# ============================================================
#  XML / FORMATTING HELPERS
# ============================================================

def remove_all_borders(cell):
    """Remove every border from a table cell."""
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    tcBorders = OxmlElement('w:tcBorders')
    for side in ['top', 'left', 'bottom', 'right', 'insideH', 'insideV']:
        border = OxmlElement(f'w:{side}')
        border.set(qn('w:val'), 'none')
        border.set(qn('w:sz'), '0')
        border.set(qn('w:space'), '0')
        border.set(qn('w:color'), 'auto')
        tcBorders.append(border)
    tcPr.append(tcBorders)


def add_top_border(cell, thickness=6, color='000000'):
    """Add a top border to a cell (used as footnote separator line)."""
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    tcBorders = OxmlElement('w:tcBorders')
    top = OxmlElement('w:top')
    top.set(qn('w:val'), 'single')
    top.set(qn('w:sz'), str(thickness))
    top.set(qn('w:space'), '0')
    top.set(qn('w:color'), color)
    tcBorders.append(top)
    tcPr.append(tcBorders)


def set_cell_width(cell, width_cm):
    """Set exact width of a table cell in cm."""
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    tcW = OxmlElement('w:tcW')
    tcW.set(qn('w:w'), str(int(width_cm * 567)))  # 567 twips per cm
    tcW.set(qn('w:type'), 'dxa')
    tcPr.append(tcW)


def set_row_height(row, height_cm, exact=True):
    """Set exact or minimum height of a table row."""
    tr = row._tr
    trPr = tr.get_or_add_trPr()
    trHeight = OxmlElement('w:trHeight')
    trHeight.set(qn('w:val'), str(int(height_cm * 567)))
    trHeight.set(qn('w:hRule'), 'exact' if exact else 'atLeast')
    trPr.append(trHeight)


def set_cell_vertical_align(cell, align='top'):
    """Set vertical text alignment within a cell."""
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    vAlign = OxmlElement('w:vAlign')
    vAlign.set(qn('w:val'), align)
    tcPr.append(vAlign)


def set_cell_padding(cell, top=0.1, bottom=0.1, left=0.15, right=0.15):
    """Set internal padding (margins) for a cell in cm."""
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    tcMar = OxmlElement('w:tcMar')
    for side, val in [('top', top), ('bottom', bottom),
                      ('left', left), ('right', right)]:
        el = OxmlElement(f'w:{side}')
        el.set(qn('w:w'), str(int(val * 567)))
        el.set(qn('w:type'), 'dxa')
        tcMar.append(el)
    tcPr.append(tcMar)


def add_paragraph_to_cell(cell, text, font_name='Georgia', font_size=11,
                           bold=False, italic=False, superscript=False,
                           space_before=0, space_after=2,
                           line_spacing=None, first_line_indent=None,
                           alignment=WD_ALIGN_PARAGRAPH.LEFT):
    """Add a styled paragraph to a table cell. Returns the paragraph."""
    para = cell.add_paragraph()
    para.alignment = alignment

    pPr = para._p.get_or_add_pPr()

    # Spacing
    pSpacing = OxmlElement('w:spacing')
    pSpacing.set(qn('w:before'), str(int(space_before * 20)))  # pt to twips
    pSpacing.set(qn('w:after'), str(int(space_after * 20)))
    if line_spacing:
        pSpacing.set(qn('w:line'), str(int(line_spacing * 240)))
        pSpacing.set(qn('w:lineRule'), 'auto')
    pPr.append(pSpacing)

    # First line indent
    if first_line_indent:
        pInd = OxmlElement('w:ind')
        pInd.set(qn('w:firstLine'), str(int(first_line_indent * 567)))
        pPr.append(pInd)

    run = para.add_run(text)
    run.font.name = font_name
    run.font.size = Pt(font_size)
    run.font.bold = bold
    run.font.italic = italic
    if superscript:
        run.font.superscript = True

    return para


def add_footnote_to_cell(cell, number, text, font_name='Georgia',
                          font_size=9, space_after=1):
    """Add a numbered footnote paragraph to a cell."""
    para = cell.add_paragraph()
    para.alignment = WD_ALIGN_PARAGRAPH.LEFT

    pPr = para._p.get_or_add_pPr()
    pSpacing = OxmlElement('w:spacing')
    pSpacing.set(qn('w:before'), '0')
    pSpacing.set(qn('w:after'), str(int(space_after * 20)))
    pPr.append(pSpacing)

    # Superscript number
    num_run = para.add_run(str(number))
    num_run.font.name = font_name
    num_run.font.size = Pt(font_size - 1)
    num_run.font.superscript = True

    # Footnote text
    text_run = para.add_run(f'\u2009{text}')  # thin space before text
    text_run.font.name = font_name
    text_run.font.size = Pt(font_size)

    return para


# ============================================================
#  DOCUMENT SETUP
# ============================================================

def create_document(page_size='A4', margin_cm=1.5):
    """
    Create a landscape Word document.
    Returns: doc, page_width_cm, page_height_cm
    """
    doc = Document()

    # Remove default empty paragraph
    for para in doc.paragraphs:
        p = para._element
        p.getparent().remove(p)

    section = doc.sections[0]

    # Landscape dimensions
    sizes = {
        'A4':     (29.7, 21.0),
        'Letter': (27.94, 21.59),
        'A5':     (21.0, 14.85),   # landscape A5
    }
    w, h = sizes.get(page_size, (29.7, 21.0))

    section.page_width  = Cm(w)
    section.page_height = Cm(h)
    section.left_margin   = Cm(margin_cm)
    section.right_margin  = Cm(margin_cm)
    section.top_margin    = Cm(margin_cm)
    section.bottom_margin = Cm(margin_cm)

    return doc, w, h


# ============================================================
#  CORE LAYOUT BUILDER
# ============================================================

def build_book_layout(doc, pages_data, page_width_cm, page_height_cm,
                      margin_cm=1.5, footnote_ratio=0.25,
                      font_name='Georgia', font_size=11,
                      footnote_font_size=9, separator_thickness=6,
                      separator_color='000000', line_spacing=1.2,
                      first_line_indent=0.5):
    """
    Build the full book layout into `doc`.

    pages_data: list of dicts
        {
          'main_text': str,
          'footnotes': [str, str, ...]   # list of footnote strings
        }

    Layout per Word page:
    ┌─────────────────┬─────────────────┐
    │   LEFT PAGE     │   RIGHT PAGE    │
    │  (main text)    │  (main text)    │
    │─────────────────│─────────────────│  ← separator line
    │  (footnotes)    │  (footnotes)    │
    └─────────────────┴─────────────────┘
    """

    # Usable dimensions
    usable_width  = page_width_cm  - (margin_cm * 2)
    usable_height = page_height_cm - (margin_cm * 2)

    gutter_cm   = 0.5
    col_width   = (usable_width - gutter_cm) / 2

    text_height = usable_height * (1 - footnote_ratio)
    foot_height = usable_height * footnote_ratio

    first_sheet = True

    # ── Pair pages: 2 per sheet ──
    for sheet_idx in range(0, len(pages_data), 2):
        pair = pages_data[sheet_idx: sheet_idx + 2]

        # Pad to always have 2 columns
        while len(pair) < 2:
            pair.append({'main_text': '', 'footnotes': []})

        # Page break between sheets (not before the first)
        if not first_sheet:
            doc.add_page_break()
        first_sheet = False

        # ── Outer table: 1 row × 2 cols ──
        outer = doc.add_table(rows=1, cols=2)
        outer.style = 'Table Grid'

        # Remove outer table borders
        tbl = outer._tbl
        tblPr = tbl.find(qn('w:tblPr'))
        if tblPr is None:
            tblPr = OxmlElement('w:tblPr')
            tbl.insert(0, tblPr)
        tblBorders = OxmlElement('w:tblBorders')
        for side in ['top', 'left', 'bottom', 'right', 'insideH', 'insideV']:
            b = OxmlElement(f'w:{side}')
            b.set(qn('w:val'), 'none')
            tblBorders.append(b)
        tblPr.append(tblBorders)

        for col_idx, (outer_cell, page) in enumerate(
                zip(outer.rows[0].cells, pair)):

            remove_all_borders(outer_cell)
            set_cell_width(outer_cell, col_width)
            set_cell_padding(outer_cell, 0, 0, 0, 0)

            # ── Inner table: 2 rows × 1 col (text | footnotes) ──
            inner = outer_cell.add_table(rows=2, cols=1)
            inner.style = 'Table Grid'

            # Remove inner table borders too
            i_tbl = inner._tbl
            i_tblPr = i_tbl.find(qn('w:tblPr'))
            if i_tblPr is None:
                i_tblPr = OxmlElement('w:tblPr')
                i_tbl.insert(0, i_tblPr)
            i_tblBorders = OxmlElement('w:tblBorders')
            for side in ['top', 'left', 'bottom', 'right',
                         'insideH', 'insideV']:
                b = OxmlElement(f'w:{side}')
                b.set(qn('w:val'), 'none')
                i_tblBorders.append(b)
            i_tblPr.append(i_tblBorders)

            text_row = inner.rows[0]
            foot_row = inner.rows[1]

            set_row_height(text_row, text_height, exact=True)
            set_row_height(foot_row, foot_height, exact=True)

            text_cell = text_row.cells[0]
            foot_cell = foot_row.cells[0]

            # Style cells
            remove_all_borders(text_cell)
            remove_all_borders(foot_cell)
            set_cell_width(text_cell, col_width)
            set_cell_width(foot_cell, col_width)
            set_cell_padding(text_cell, 0.1, 0.1, 0.2, 0.2)
            set_cell_padding(foot_cell, 0.15, 0.1, 0.2, 0.2)
            set_cell_vertical_align(text_cell, 'top')
            set_cell_vertical_align(foot_cell, 'top')

            # Separator line above footnote area
            add_top_border(foot_cell,
                           thickness=separator_thickness,
                           color=separator_color)

            # ── Main text ──
            if page['main_text'].strip():
                # Split by newlines to preserve paragraph breaks
                paragraphs = page['main_text'].split('\n')
                for p_idx, para_text in enumerate(paragraphs):
                    if para_text.strip():
                        add_paragraph_to_cell(
                            text_cell,
                            para_text,
                            font_name=font_name,
                            font_size=font_size,
                            space_before=0,
                            space_after=3,
                            line_spacing=line_spacing,
                            first_line_indent=(
                                first_line_indent if p_idx > 0 else 0
                            )
                        )

            # ── Footnotes ──
            has_footnotes = any(
                f.strip() for f in page.get('footnotes', [])
            )
            if has_footnotes:
                fn_number = 1
                for fn_text in page['footnotes']:
                    if fn_text.strip():
                        add_footnote_to_cell(
                            foot_cell,
                            fn_number,
                            fn_text,
                            font_name=font_name,
                            font_size=footnote_font_size
                        )
                        fn_number += 1

    return doc


# ============================================================
#  STREAMLIT UI
# ============================================================

def main():
    st.set_page_config(
        page_title='Book Layout Generator',
        page_icon='📖',
        layout='wide'
    )

    # ── Header ──
    st.title('📖 Book Layout Generator')
    st.caption(
        'Creates a landscape Word document where each printed sheet = '
        '2 book pages. Print double‑sided → cut vertically → stack as a book.'
    )

    # ── Sidebar ──
    with st.sidebar:
        st.header('⚙️ Layout Settings')

        page_size = st.selectbox(
            'Paper Size',
            ['A4', 'Letter', 'A5'],
            index=0
        )

        margin_cm = st.slider(
            'Page Margin (cm)', 0.5, 3.0, 1.5, 0.1
        )

        footnote_ratio = st.slider(
            'Footnote Area — % of page height',
            5, 70, 25, 5,
            help='How much of each half-page is reserved for footnotes.'
        ) / 100

        st.divider()
        st.subheader('🔤 Typography')

       font_name = st.selectbox(
    'Font',
    [
        'Courier New',
        'Consolas',
        'Arial',
        'Times New Roman',
        'Courier',
        'EB Garamond',        # free / open source
        'Cormorant Garamond', # free / open source
        'UnifrakturCook',     # gothic / blackletter - free
        'MedievalSharp',      # gothic - free
        'Cinzel',             # classical - free
    ],
    index=0
)

        font_size = st.slider('Main Text (pt)', 8, 16, 11)

        footnote_font_size = st.slider('Footnote Text (pt)', 6, 13, 9)

        line_spacing = st.slider(
            'Line Spacing', 1.0, 2.0, 1.2, 0.1
        )

        first_line_indent = st.slider(
            'First Line Indent (cm)', 0.0, 1.5, 0.5, 0.1
        )

        st.divider()
        st.subheader('📏 Separator Line')

        sep_thickness = st.slider(
            'Separator Thickness (pt)', 1, 12, 6
        )

        sep_color = st.color_picker(
            'Separator Color', '#000000'
        ).lstrip('#')

        st.divider()
        st.info(
            '**How to print:**\n\n'
            '1. Print double‑sided (flip on long edge)\n'
            '2. Cut each sheet down the vertical centre\n'
            '3. Stack halves in page order\n'
            '4. Bind or fold'
        )

    # ── Session state init ──
    if 'pages' not in st.session_state:
        st.session_state.pages = [
            {
                'main_text': 'Your text goes here. You can type or paste '
                             'multiple paragraphs.\n\nThis is a second '
                             'paragraph.',
                'footnotes': ['This is a sample footnote.']
            }
        ]

    # ── Page list controls ──
    st.subheader('📄 Pages')
    st.caption(
        'Each entry below = one half‑page (one book page). '
        'Pages are paired left–right on each printed sheet.'
    )

    col_add, col_clear = st.columns([1, 1])
    with col_add:
        if st.button('➕ Add Page', use_container_width=True):
            st.session_state.pages.append(
                {'main_text': '', 'footnotes': []}
            )
            st.rerun()
    with col_clear:
        if st.button('🗑 Clear All Pages', use_container_width=True):
            st.session_state.pages = [
                {'main_text': '', 'footnotes': []}
            ]
            st.rerun()

    st.divider()

    # ── Page editors ──
    pages_to_delete = []

    for i, page in enumerate(st.session_state.pages):
        with st.expander(
            f'📃 Page {i + 1}   '
            f'{"(Sheet " + str((i // 2) + 1) + ", " + ("Left" if i % 2 == 0 else "Right") + " column)"}',
            expanded=(i < 2)
        ):
            left_col, right_col = st.columns([3, 2])

            # ── Main text ──
            with left_col:
                st.markdown('**Main Text**')
                st.caption(
                    'Use blank lines to separate paragraphs.'
                )
                page['main_text'] = st.text_area(
                    f'main_text_{i}',
                    value=page['main_text'],
                    height=220,
                    key=f'main_{i}',
                    label_visibility='collapsed'
                )

            # ── Footnotes ──
            with right_col:
                st.markdown('**Footnotes**')
                st.caption(
                    'Numbered automatically starting from 1 per page.'
                )

                updated_footnotes = []
                fn_to_delete = None

                for j, fn in enumerate(page.get('footnotes', [])):
                    fn_input_col, fn_del_col = st.columns([5, 1])
                    with fn_input_col:
                        updated_fn = st.text_input(
                            f'Footnote {j + 1}',
                            value=fn,
                            key=f'fn_{i}_{j}'
                        )
                        updated_footnotes.append(updated_fn)
                    with fn_del_col:
                        st.markdown('<br>', unsafe_allow_html=True)
                        if st.button('✕', key=f'del_fn_{i}_{j}',
                                     help='Remove this footnote'):
                            fn_to_delete = j

                # Apply footnote deletion
                if fn_to_delete is not None:
                    updated_footnotes.pop(fn_to_delete)
                    page['footnotes'] = updated_footnotes
                    st.rerun()
                else:
                    page['footnotes'] = updated_footnotes

                if st.button(
                    '＋ Add Footnote',
                    key=f'add_fn_{i}',
                    use_container_width=True
                ):
                    page['footnotes'].append('')
                    st.rerun()

            # ── Remove page ──
            st.divider()
            if st.button(
                f'🗑 Remove Page {i + 1}',
                key=f'del_page_{i}',
                use_container_width=False
            ):
                if len(st.session_state.pages) > 1:
                    pages_to_delete.append(i)

    # Apply page deletions
    if pages_to_delete:
        for idx in sorted(pages_to_delete, reverse=True):
            st.session_state.pages.pop(idx)
        st.rerun()

    # ── Sheet preview ──
    st.divider()
    st.subheader('🗂 Sheet Overview')
    total_pages  = len(st.session_state.pages)
    total_sheets = (total_pages + 1) // 2
    st.write(
        f'**{total_pages} pages** → '
        f'**{total_sheets} printed sheets** '
        f'(double-sided)'
    )

    cols = st.columns(min(total_sheets, 4))
    for s in range(total_sheets):
        left_page  = s * 2
        right_page = s * 2 + 1
        with cols[s % 4]:
            right_label = (
                f'Page {right_page + 1}'
                if right_page < total_pages
                else '(blank)'
            )
            st.info(
                f'**Sheet {s + 1}**\n\n'
                f'Left → Page {left_page + 1}\n\n'
                f'Right → {right_label}'
            )

    # ── Generate button ──
    st.divider()
    generate_col, _ = st.columns([2, 3])
    with generate_col:
        generate = st.button(
            '📥 Generate Word Document',
            type='primary',
            use_container_width=True
        )

    if generate:
        with st.spinner('Building your document...'):
            try:
                doc, pw, ph = create_document(page_size, margin_cm)

                # Clean up empty footnotes
                clean_pages = []
                for p in st.session_state.pages:
                    clean_pages.append({
                        'main_text': p.get('main_text', ''),
                        'footnotes': [
                            f for f in p.get('footnotes', [])
                            if f.strip()
                        ]
                    })

                doc = build_book_layout(
                    doc=doc,
                    pages_data=clean_pages,
                    page_width_cm=pw,
                    page_height_cm=ph,
                    margin_cm=margin_cm,
                    footnote_ratio=footnote_ratio,
                    font_name=font_name,
                    font_size=font_size,
                    footnote_font_size=footnote_font_size,
                    separator_thickness=sep_thickness,
                    separator_color=sep_color,
                    line_spacing=line_spacing,
                    first_line_indent=first_line_indent
                )

                buffer = io.BytesIO()
                doc.save(buffer)
                buffer.seek(0)

                st.success('✅ Document ready to download!')
                st.download_button(
                    label='⬇️ Download book_layout.docx',
                    data=buffer,
                    file_name='book_layout.docx',
                    mime=(
                        'application/vnd.openxmlformats-officedocument'
                        '.wordprocessingml.document'
                    ),
                    use_container_width=True
                )

            except Exception as e:
                st.error(f'Something went wrong: {e}')
                st.exception(e)


# ============================================================
#  ENTRY POINT
# ============================================================

if __name__ == '__main__':
    main()
